"""Build complex Word (docx) documents from structured JSON definitions.

Supports headings, paragraphs with inline formatting, tables,
images, page breaks, headers/footers, and bullet/numbered lists.
"""

from __future__ import annotations

import base64
import io
from typing import Any, ClassVar, Literal

from pydantic import BaseModel, Field

from llming_plumber.blocks.base import BaseBlock, BlockContext, BlockInput, BlockOutput


class RunDef(BaseModel):
    """A run of text with uniform formatting within a paragraph."""

    text: str = Field(description="Text content")
    bold: bool = False
    italic: bool = False
    underline: bool = False
    strike: bool = False
    font_name: str = Field(default="", description="Font name, e.g. 'Arial'")
    font_size: int | None = Field(
        default=None, description="Font size in points"
    )
    font_color: str = Field(
        default="",
        description="Hex color without #, e.g. 'FF0000'",
    )
    highlight: str = Field(
        default="",
        description="Highlight color name: yellow, green, red, blue, etc.",
    )


class ParagraphDef(BaseModel):
    """A paragraph with optional inline formatting runs."""

    text: str = Field(
        default="",
        description="Simple text. Ignored if runs are provided.",
    )
    runs: list[RunDef] = Field(
        default=[],
        description="Formatted text runs. Overrides text if provided.",
    )
    style: str = Field(
        default="Normal",
        description=(
            "Word paragraph style: Normal, Heading 1, Heading 2, Heading 3, "
            "List Bullet, List Number, Quote, Title, Subtitle"
        ),
    )
    alignment: Literal["left", "center", "right", "justify", ""] = ""
    space_before: int | None = Field(
        default=None, description="Space before paragraph in points"
    )
    space_after: int | None = Field(
        default=None, description="Space after paragraph in points"
    )


class TableCellDef(BaseModel):
    """A single cell in a table."""

    text: str = Field(default="", description="Cell text content")
    bold: bool = False
    bg_color: str = Field(
        default="",
        description="Cell background hex color without #",
    )
    colspan: int = Field(default=1, description="Column span")


class TableRowDef(BaseModel):
    """A row of cells in a table."""

    cells: list[TableCellDef]
    is_header: bool = Field(
        default=False,
        description="Mark as header row (repeats on page breaks)",
    )


class DocTableDef(BaseModel):
    """A table in the document."""

    rows: list[TableRowDef]
    style: str = Field(
        default="Table Grid",
        description="Word table style name",
    )
    col_widths: list[float] = Field(
        default=[],
        description="Column widths in inches. Empty for auto.",
    )


class ImageDef(BaseModel):
    """An inline image."""

    data_b64: str = Field(description="Base64-encoded image (PNG or JPEG)")
    width_inches: float = Field(
        default=0, description="Image width in inches. 0 for original."
    )
    height_inches: float = Field(
        default=0, description="Image height in inches. 0 for original."
    )
    caption: str = Field(default="", description="Optional caption below image")


class SectionDef(BaseModel):
    """A section of the document with its own layout settings."""

    orientation: Literal["portrait", "landscape"] = "portrait"
    elements: list[dict[str, Any]] = Field(
        description=(
            "Ordered list of elements. Each dict has a 'type' key: "
            "'paragraph', 'table', 'image', 'page_break'. "
            "Remaining keys match the corresponding definition model."
        ),
    )
    header_text: str = Field(
        default="", description="Header text for this section"
    )
    footer_text: str = Field(
        default="", description="Footer text for this section"
    )


class WordBuilderInput(BlockInput):
    sections: list[SectionDef] = Field(
        title="Sections",
        description="Document sections with elements",
        min_length=1,
    )
    default_font: str = Field(
        default="Calibri",
        title="Default Font",
        description="Default font for the document",
    )
    default_font_size: int = Field(
        default=11,
        title="Default Font Size",
        description="Default font size in points",
    )
    json_data: str = Field(
        default="",
        title="JSON Data",
        description=(
            "Alternative: provide sections as JSON matching SectionDef. "
            "Ignored if sections has data."
        ),
        json_schema_extra={"widget": "code", "rows": 20},
    )


class WordBuilderOutput(BlockOutput):
    content: str = Field(description="Base64-encoded docx file")
    section_count: int
    element_count: int


class WordBuilderBlock(BaseBlock[WordBuilderInput, WordBuilderOutput]):
    block_type: ClassVar[str] = "word_builder"
    icon: ClassVar[str] = "tabler/file-text"
    categories: ClassVar[list[str]] = ["documents", "word"]
    description: ClassVar[str] = (
        "Build complex Word documents with headings, tables, "
        "images, inline formatting, and multiple sections"
    )
    cache_ttl: ClassVar[int] = 0

    async def execute(
        self, input: WordBuilderInput, ctx: BlockContext | None = None
    ) -> WordBuilderOutput:
        import json

        sections = input.sections
        if input.json_data and not any(s.elements for s in sections):
            raw = json.loads(input.json_data)
            sections = [SectionDef.model_validate(s) for s in raw]

        return self._build_document(
            sections, input.default_font, input.default_font_size
        )

    def _build_document(
        self,
        sections: list[SectionDef],
        default_font: str,
        default_font_size: int,
    ) -> WordBuilderOutput:
        import docx
        from docx.shared import Inches, Pt, RGBColor

        doc = docx.Document()

        style = doc.styles["Normal"]
        style.font.name = default_font
        style.font.size = Pt(default_font_size)

        total_elements = 0

        for sec_idx, section in enumerate(sections):
            if sec_idx > 0:
                doc.add_page_break()

            if section.header_text:
                header = doc.sections[-1].header
                header.is_linked_to_previous = False
                hp = header.paragraphs[0]
                hp.text = section.header_text

            if section.footer_text:
                footer = doc.sections[-1].footer
                footer.is_linked_to_previous = False
                fp = footer.paragraphs[0]
                fp.text = section.footer_text

            for elem in section.elements:
                elem_type = elem.get("type", "paragraph")

                if elem_type == "paragraph":
                    self._add_paragraph(doc, elem, Pt, RGBColor)
                elif elem_type == "table":
                    self._add_table(doc, elem, Pt, RGBColor)
                elif elem_type == "image":
                    self._add_image(doc, elem, Inches)
                elif elem_type == "page_break":
                    doc.add_page_break()

                total_elements += 1

        buf = io.BytesIO()
        doc.save(buf)
        encoded = base64.b64encode(buf.getvalue()).decode("ascii")

        return WordBuilderOutput(
            content=encoded,
            section_count=len(sections),
            element_count=total_elements,
        )

    def _add_paragraph(
        self, doc: Any, elem: dict[str, Any], Pt: Any, RGBColor: Any
    ) -> None:
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        para_def = ParagraphDef.model_validate(
            {k: v for k, v in elem.items() if k != "type"}
        )
        para = doc.add_paragraph(style=para_def.style)

        if para_def.alignment:
            align_map = {
                "left": WD_ALIGN_PARAGRAPH.LEFT,
                "center": WD_ALIGN_PARAGRAPH.CENTER,
                "right": WD_ALIGN_PARAGRAPH.RIGHT,
                "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
            }
            para.alignment = align_map.get(para_def.alignment)

        if para_def.space_before is not None:
            para.paragraph_format.space_before = Pt(para_def.space_before)
        if para_def.space_after is not None:
            para.paragraph_format.space_after = Pt(para_def.space_after)

        if para_def.runs:
            for run_def in para_def.runs:
                run = para.add_run(run_def.text)
                run.bold = run_def.bold
                run.italic = run_def.italic
                run.underline = run_def.underline
                if run_def.strike:
                    run.font.strike = True
                if run_def.font_name:
                    run.font.name = run_def.font_name
                if run_def.font_size:
                    run.font.size = Pt(run_def.font_size)
                if run_def.font_color:
                    h = run_def.font_color
                    run.font.color.rgb = RGBColor(
                        int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16)
                    )
        else:
            para.add_run(para_def.text)

    def _add_table(
        self, doc: Any, elem: dict[str, Any], Pt: Any, RGBColor: Any
    ) -> None:
        from docx.oxml.ns import qn

        table_def = DocTableDef.model_validate(
            {k: v for k, v in elem.items() if k != "type"}
        )
        if not table_def.rows:
            return

        num_cols = max(len(r.cells) for r in table_def.rows)
        table = doc.add_table(
            rows=len(table_def.rows), cols=num_cols
        )
        table.style = table_def.style

        for row_idx, row_def in enumerate(table_def.rows):
            row = table.rows[row_idx]
            for col_idx, cell_def in enumerate(row_def.cells):
                if col_idx >= num_cols:
                    break
                cell = row.cells[col_idx]
                cell.text = cell_def.text

                if cell_def.bold:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.bold = True

                if cell_def.bg_color:
                    shading = cell._element.get_or_add_tcPr()
                    shading_el = shading.makeelement(
                        qn("w:shd"), {
                            qn("w:fill"): cell_def.bg_color,
                            qn("w:val"): "clear",
                        }
                    )
                    shading.append(shading_el)

    def _add_image(
        self, doc: Any, elem: dict[str, Any], Inches: Any
    ) -> None:
        img_def = ImageDef.model_validate(
            {k: v for k, v in elem.items() if k != "type"}
        )
        img_data = base64.b64decode(img_def.data_b64)

        kwargs: dict[str, Any] = {}
        if img_def.width_inches > 0:
            kwargs["width"] = Inches(img_def.width_inches)
        if img_def.height_inches > 0:
            kwargs["height"] = Inches(img_def.height_inches)

        doc.add_picture(io.BytesIO(img_data), **kwargs)

        if img_def.caption:
            doc.add_paragraph(img_def.caption, style="Caption")
