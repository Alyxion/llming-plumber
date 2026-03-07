"""Read Word (docx) documents."""

from __future__ import annotations

import base64
import io
from typing import Any, ClassVar

from pydantic import Field

from llming_plumber.blocks.base import BaseBlock, BlockContext, BlockInput, BlockOutput
from llming_plumber.blocks.limits import check_base64_size, check_file_size


class WordReaderInput(BlockInput):
    content: str = Field(
        title="Content",
        description="Base64-encoded docx file bytes",
        json_schema_extra={"widget": "textarea"},
    )


class WordReaderOutput(BlockOutput):
    text: str
    paragraphs: list[dict[str, str]]
    tables: list[list[list[str]]]
    metadata: dict[str, Any]


class WordReaderBlock(BaseBlock[WordReaderInput, WordReaderOutput]):
    block_type: ClassVar[str] = "word_reader"
    icon: ClassVar[str] = "tabler/file-text"
    categories: ClassVar[list[str]] = ["documents", "word"]
    description: ClassVar[str] = "Read Word (docx) documents"
    cache_ttl: ClassVar[int] = 0

    async def execute(
        self, input: WordReaderInput, ctx: BlockContext | None = None
    ) -> WordReaderOutput:
        import docx

        check_base64_size(input.content, label="Word file")
        raw = base64.b64decode(input.content)
        check_file_size(len(raw), label="Word file")
        doc = docx.Document(io.BytesIO(raw))

        paragraphs: list[dict[str, str]] = []
        text_parts: list[str] = []
        for para in doc.paragraphs:
            paragraphs.append({
                "text": para.text,
                "style": para.style.name if para.style else "Normal",
            })
            text_parts.append(para.text)

        tables: list[list[list[str]]] = []
        for table in doc.tables:
            table_data: list[list[str]] = []
            for row in table.rows:
                table_data.append([cell.text for cell in row.cells])
            tables.append(table_data)

        metadata: dict[str, Any] = {}
        props = doc.core_properties
        for attr in ("author", "title", "subject", "keywords", "category", "comments"):
            val = getattr(props, attr, None)
            if val:
                metadata[attr] = str(val)
        if props.created:
            metadata["created"] = props.created.isoformat()
        if props.modified:
            metadata["modified"] = props.modified.isoformat()

        return WordReaderOutput(
            text="\n".join(text_parts),
            paragraphs=paragraphs,
            tables=tables,
            metadata=metadata,
        )
