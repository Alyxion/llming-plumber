"""Extract structured SectionDef models from Word documents.

Outputs the same Pydantic models that WordBuilderBlock consumes,
enabling round-trip: extract → edit JSON → rebuild.
"""

from __future__ import annotations

import base64
import io
from typing import Any, ClassVar

from pydantic import Field

from llming_plumber.blocks.base import (
    BaseBlock,
    BlockContext,
    BlockInput,
    BlockOutput,
)
from llming_plumber.blocks.documents.word_builder import (
    SectionDef,
)
from llming_plumber.blocks.limits import check_base64_size, check_file_size


class WordExtractorInput(BlockInput):
    content: str = Field(
        title="Content",
        description="Base64-encoded docx file bytes",
        json_schema_extra={"widget": "textarea"},
    )
    extract_formatting: bool = Field(
        default=True,
        title="Extract Formatting",
        description="Include inline run formatting (bold, italic, etc.)",
    )


class WordExtractorOutput(BlockOutput):
    sections: list[SectionDef]
    sections_json: str = Field(
        description=(
            "JSON string of sections for piping to builders"
        ),
    )


class WordExtractorBlock(
    BaseBlock[WordExtractorInput, WordExtractorOutput]
):
    block_type: ClassVar[str] = "word_extractor"
    icon: ClassVar[str] = "tabler/file-text"
    categories: ClassVar[list[str]] = ["documents", "word"]
    description: ClassVar[str] = (
        "Extract structured SectionDef models from Word documents "
        "for editing and rebuilding"
    )
    cache_ttl: ClassVar[int] = 0

    async def execute(
        self,
        input: WordExtractorInput,
        ctx: BlockContext | None = None,
    ) -> WordExtractorOutput:
        import json

        import docx

        check_base64_size(input.content, label="Word file")
        raw = base64.b64decode(input.content)
        check_file_size(len(raw), label="Word file")
        doc = docx.Document(io.BytesIO(raw))

        elements: list[dict[str, Any]] = []

        for para in doc.paragraphs:
            elem = self._extract_paragraph(
                para, input.extract_formatting
            )
            if elem:
                elements.append(elem)

        for table in doc.tables:
            elements.append(self._extract_table(table))

        header_text = ""
        footer_text = ""
        if doc.sections:
            sec = doc.sections[0]
            if sec.header and sec.header.paragraphs:
                header_text = sec.header.paragraphs[0].text
            if sec.footer and sec.footer.paragraphs:
                footer_text = sec.footer.paragraphs[0].text

        section = SectionDef(
            elements=elements,
            header_text=header_text,
            footer_text=footer_text,
        )

        sections_json = json.dumps(
            [section.model_dump()], default=str
        )

        return WordExtractorOutput(
            sections=[section],
            sections_json=sections_json,
        )

    def _extract_paragraph(
        self, para: Any, extract_formatting: bool
    ) -> dict[str, Any] | None:
        if not para.text.strip():
            return None

        style = (
            para.style.name if para.style else "Normal"
        )
        elem: dict[str, Any] = {"type": "paragraph"}

        if extract_formatting and para.runs:
            runs = []
            for run in para.runs:
                run_dict: dict[str, Any] = {
                    "text": run.text,
                }
                if run.bold:
                    run_dict["bold"] = True
                if run.italic:
                    run_dict["italic"] = True
                if run.underline:
                    run_dict["underline"] = True
                if run.font.strike:
                    run_dict["strike"] = True
                if run.font.name:
                    run_dict["font_name"] = run.font.name
                if run.font.size:
                    run_dict["font_size"] = (
                        run.font.size.pt
                        if hasattr(run.font.size, "pt")
                        else run.font.size
                    )
                if (
                    run.font.color
                    and run.font.color.rgb
                ):
                    run_dict["font_color"] = str(
                        run.font.color.rgb
                    )
                runs.append(run_dict)

            if any(
                len(r) > 1 for r in runs
            ):
                elem["runs"] = runs
            else:
                elem["text"] = para.text
        else:
            elem["text"] = para.text

        elem["style"] = style
        return elem

    def _extract_table(
        self, table: Any
    ) -> dict[str, Any]:
        rows = []
        for row in table.rows:
            cells = []
            for cell in row.cells:
                cells.append({"text": cell.text})
            rows.append({"cells": cells})
        return {"type": "table", "rows": rows}
