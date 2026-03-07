"""Write Word (docx) documents."""

from __future__ import annotations

import base64
import io
from typing import ClassVar

from pydantic import Field

from llming_plumber.blocks.base import BaseBlock, BlockContext, BlockInput, BlockOutput


class WordWriterInput(BlockInput):
    paragraphs: list[dict[str, str]] = Field(
        title="Paragraphs",
        description=(
            "List of paragraph dicts with 'text' and optional 'style' "
            "(Heading1, Heading2, Normal, ListBullet)"
        ),
    )
    title: str = Field(
        default="",
        title="Title",
        description="Optional document title (added as first heading)",
    )


class WordWriterOutput(BlockOutput):
    content: str
    paragraph_count: int


class WordWriterBlock(BaseBlock[WordWriterInput, WordWriterOutput]):
    block_type: ClassVar[str] = "word_writer"
    icon: ClassVar[str] = "tabler/file-text"
    categories: ClassVar[list[str]] = ["documents", "word"]
    description: ClassVar[str] = "Write Word (docx) documents"
    cache_ttl: ClassVar[int] = 0

    async def execute(
        self, input: WordWriterInput, ctx: BlockContext | None = None
    ) -> WordWriterOutput:
        import docx

        doc = docx.Document()

        if input.title:
            doc.add_heading(input.title, level=0)

        for para in input.paragraphs:
            text = para.get("text", "")
            style = para.get("style", "Normal")
            doc.add_paragraph(text, style=style)

        buf = io.BytesIO()
        doc.save(buf)
        encoded = base64.b64encode(buf.getvalue()).decode("ascii")

        return WordWriterOutput(
            content=encoded,
            paragraph_count=len(input.paragraphs),
        )
